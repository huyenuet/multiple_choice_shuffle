import copy
from docx import Document
from engine.ParagraphList import Paragraph_List
from engine.Question import Question
import sys

def get_text(doc_name):
    full_text = []
    for para in doc_name.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


# use swap if: original doc already has paragraph style
def swap(doc_name):
    list_question = []
    answer_length = 0
    question = Question()
    for para in doc_name.paragraphs:
        if para.style.name == 'List Number':
            question.text = para.text
        elif para.style.name == 'List Number 2':
            question.add_answer(para.text)
            answer_length += 1
        if answer_length == 4:
            question.shuffle_answer()
            list_question.append(copy.deepcopy(question))
            question.reset()
            answer_length = 0
    return list_question


# use swap if: original doc already has paragraph style
def swap2(doc_name):
    list_question = []
    answer_length = 0
    question = Question()
    for para in doc_name.paragraphs:
        end_index = para.text.find(".")
        order_text = para.text[0:end_index].strip()
        if order_text.isdigit():
            question.text = para.text[end_index+1:len(para.text)].strip()
        else:
            question.add_answer(para.text[end_index+1:len(para.text)].strip())
            answer_length += 1
        if answer_length == 4:
            question.shuffle_answer()
            list_question.append(copy.deepcopy(question))
            question.reset()
            answer_length = 0
    return list_question


def create_new_doc(list_question, my_list):
    for question in list_question:
        my_list.add_item(question.text, 1)
        for answer in question.answers:
            my_list.add_item(answer.text, 2)


# doc = Document("de_kiem_tra.docx")
#
# # shuffle questions and return a list of questions
# question_list = swap2(doc)
#
# # initial a new doc
# new_doc = Document()
#
# # create an ordered list
# my_list = Paragraph_List(new_doc)
# create_new_doc(question_list, my_list)
# new_doc.save('new.docx')


file_input = sys.argv[1]

doc = Document(file_input)
question_list = swap2(doc)
new_doc = Document()
my_list = Paragraph_List(new_doc)
create_new_doc(question_list, my_list)
new_doc.save('new.docx')

sys.stdout.flush()
